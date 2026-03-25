from typing import Annotated, TypedDict

from dotenv import load_dotenv
from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import END, START, StateGraph
from langgraph.graph.message import add_messages
from langchain_groq import ChatGroq


class AgentState(TypedDict):
    messages: Annotated[list, add_messages]
    markdown_content: str
    docx_path: str


def writer_node(state: AgentState) -> AgentState:
    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.4)

    system_prompt = (
        "Actua como un redactor experto. Devuelve el contenido en Markdown estricto, "
        "usando '# ' para el titulo principal, '## ' para subtitulos y texto normal "
        "para los parrafos. Puedes usar listas con '- ' o '* ' si ayudan. No uses "
        "bloques de codigo ni fences como ```markdown. Devuelve solo el texto crudo."
    )

    user_prompt = state["messages"][-1].content
    response = llm.invoke(
        [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt),
        ]
    )

    return {
        "messages": [response],
        "markdown_content": response.content,
    }


def docx_generator_node(state: AgentState) -> AgentState:
    doc = Document()

    for raw_line in state["markdown_content"].splitlines():
        line = raw_line.strip()

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line:
            doc.add_paragraph(line)

    output_path = "salida.docx"
    doc.save(output_path)

    return {"docx_path": output_path}


graph_builder = StateGraph(AgentState)
graph_builder.add_node("writer", writer_node)
graph_builder.add_node("generator", docx_generator_node)
graph_builder.add_edge(START, "writer")
graph_builder.add_edge("writer", "generator")
graph_builder.add_edge("generator", END)
graph = graph_builder.compile()


if __name__ == "__main__":
    load_dotenv()

    prompt = (
        "Genera una propuesta comercial breve para un servicio de consultoria IA, "
        "con introduccion, servicios y precios"
    )

    result = graph.invoke(
        {
            "messages": [HumanMessage(content=prompt)],
            "markdown_content": "",
            "docx_path": "",
        }
    )

    print(f"Documento generado en: {result['docx_path']}")
